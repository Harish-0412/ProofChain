"""Generate the deterministic three-department ProofChain mock institution."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import shutil
from dataclasses import asdict, dataclass
from datetime import date
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ACADEMIC_YEAR = "2025-2026"
REQUIREMENTS = {
    "C3.2.1": "Industry interaction activities",
    "C5.1.3": "Student enrichment programmes",
    "C6.3.2": "Faculty development activities",
    "C7.1.1": "Extension and outreach activities",
    "C1.2.1": "Value-added courses",
}
DEPARTMENTS = {
    "AIML": {
        "name": "Artificial Intelligence and Machine Learning",
        "coordinator": "Dr. Ananya Rao",
        "program": "B.Tech Artificial Intelligence and Machine Learning",
    },
    "AIDS": {
        "name": "Artificial Intelligence and Data Science",
        "coordinator": "Dr. Rohan Iyer",
        "program": "B.Tech Artificial Intelligence and Data Science",
    },
    "CSE": {
        "name": "Computer Science and Engineering",
        "coordinator": "Dr. Meera Nair",
        "program": "B.Tech Computer Science and Engineering",
    },
}
EVENT_TITLES = {
    "AIML": {
        "C3.2.1": "Industry Workshop on Production Machine Learning",
        "C5.1.3": "Applied Machine Learning Skills Bootcamp",
        "C6.3.2": "Faculty Development Programme on Responsible AI",
        "C7.1.1": "Community AI Literacy Outreach",
        "C1.2.1": "MLOps Value-Added Certification Course",
    },
    "AIDS": {
        "C3.2.1": "Industry Workshop on Modern Data Engineering",
        "C5.1.3": "Analytics and Visualization Skills Bootcamp",
        "C6.3.2": "Faculty Development Programme on Data Governance",
        "C7.1.1": "Community Data Literacy Outreach",
        "C1.2.1": "Business Intelligence Value-Added Course",
    },
    "CSE": {
        "C3.2.1": "Cloud Native Engineering Industry Workshop",
        "C5.1.3": "Secure Coding Student Enrichment Bootcamp",
        "C6.3.2": "Faculty Development Programme on Software Architecture",
        "C7.1.1": "Community Cyber Safety Outreach",
        "C1.2.1": "Full Stack Development Value-Added Course",
    },
}
EVENT_DATES = [
    date(2025, 8, 22),
    date(2025, 10, 17),
    date(2025, 12, 5),
    date(2026, 1, 23),
    date(2026, 3, 13),
]
EVENT_REPORT_FORMATS = {
    "AIML": ["pdf", "docx", "html", "md", "json"],
    "AIDS": ["xml", "txt", "pdf", "docx", "html"],
    "CSE": ["md", "json", "xml", "txt", "pdf"],
}
ATTENDANCE_FORMATS = {
    "AIML": ["xlsx", "csv", "tsv", "csv", "tsv"],
    "AIDS": ["xlsx", "tsv", "csv", "tsv", "csv"],
    "CSE": ["xlsx", "csv", "tsv", "csv", "tsv"],
}
APPROVAL_FORMATS = {
    "AIML": ["pdf", "docx", "txt", "json", "xml"],
    "AIDS": ["html", "pdf", "docx", "txt", "json"],
    "CSE": ["xml", "html", "pdf", "docx", "txt"],
}
CERTIFICATE_FORMATS = {
    "AIML": ["pdf", "docx", "html", "md", "txt"],
    "AIDS": ["json", "xml", "pdf", "docx", "html"],
    "CSE": ["md", "txt", "json", "xml", "pdf"],
}

FEMALE_NAMES = [
    "Aadhya Menon", "Aarohi Shah", "Akshara Pillai", "Anika Reddy", "Anvi Joshi",
    "Avni Kulkarni", "Diya Nair", "Eesha Kapoor", "Harini Iyer", "Ishita Rao",
    "Kavya Bhat", "Meera Das", "Nandini Bose", "Riya Verma", "Saanvi Gupta",
    "Aisha Khan", "Bhavya Jain", "Charvi Shetty", "Deepika Suresh", "Gauri Patil",
    "Janani Krishnan", "Keerthana Ravi", "Lakshmi Prasad", "Navya Kumar", "Nithya Sen",
    "Pavithra Anand", "Roshni Thomas", "Shruti Desai", "Tanvi Mishra", "Veda Narayan",
    "Amrita George", "Deeksha Singh", "Hema Chandran", "Ira Banerjee", "Jahnavi Naidu",
    "Kriti Malhotra", "Madhavi Gopal", "Neha Srinivas", "Pooja Saxena", "Radhika Roy",
    "Sanjana Arora", "Sneha Varma", "Swathi Ramesh", "Trisha Mehta", "Yamini Sethi",
]
MALE_NAMES = [
    "Aarav Sharma", "Aditya Nair", "Akash Reddy", "Arjun Menon", "Atharv Joshi",
    "Dev Patel", "Dhruv Iyer", "Harsh Verma", "Ishaan Rao", "Karan Shah",
    "Madhav Bhat", "Nikhil Das", "Rohan Kumar", "Sai Prasad", "Vivaan Gupta",
    "Abhinav Suresh", "Anirudh Jain", "Arnav Shetty", "Darshan Pillai", "Gautam Sen",
    "Karthik Ravi", "Manav Bose", "Neeraj Anand", "Pranav Patil", "Rahul Thomas",
    "Ritvik Krishnan", "Siddharth Desai", "Tejas Kulkarni", "Varun Mishra", "Yash Singh",
    "Ajay George", "Deepak Chandran", "Hari Narayan", "Jay Mehta", "Krish Malhotra",
    "Mohit Naidu", "Naveen Gopal", "Raj Saxena", "Rishi Roy", "Sanjay Arora",
    "Shrey Ramesh", "Suraj Banerjee", "Vijay Kapoor", "Vikram Sethi", "Zayan Khan",
]


@dataclass(frozen=True)
class Student:
    roll_number: str
    full_name: str
    gender: str
    department: str
    program: str
    semester: int
    batch: str
    email: str
    active: bool = True


def slug(value: str) -> str:
    return "".join(character.lower() if character.isalnum() else "." for character in value).strip(".")


def build_students(department: str, department_index: int) -> list[Student]:
    definition = DEPARTMENTS[department]
    female_start = department_index * 15
    male_start = department_index * 15
    students: list[Student] = []
    names = [
        *((name, "Female") for name in FEMALE_NAMES[female_start : female_start + 15]),
        *((name, "Male") for name in MALE_NAMES[male_start : male_start + 15]),
    ]
    for sequence, (name, gender) in enumerate(names, start=1):
        roll_number = f"25{department}{sequence:03d}"
        students.append(
            Student(
                roll_number=roll_number,
                full_name=name,
                gender=gender,
                department=department,
                program=definition["program"],
                semester=2,
                batch="2025-2029",
                email=f"{slug(name)}.{department.lower()}{sequence:02d}@students.example.invalid",
            )
        )
    return students


def event_fields(department: str, requirement: str, index: int) -> dict[str, Any]:
    event_id = f"EVT-{department}-{index + 1:03d}"
    event_date = EVENT_DATES[index].isoformat()
    return {
        "Event ID": event_id,
        "Accreditation Requirement": requirement,
        "Mapped Requirement": requirement,
        "Department": department,
        "Academic Year": ACADEMIC_YEAR,
        "Event Title": EVENT_TITLES[department][requirement],
        "Event Date": event_date,
        "Coordinator": DEPARTMENTS[department]["coordinator"],
        "Reported Participant Count": 30,
    }


def labelled_lines(fields: dict[str, Any]) -> list[str]:
    return [f"{key} | {value}" for key, value in fields.items()]


def write_plain(path: Path, fields: dict[str, Any], heading: str, narrative: str) -> None:
    content = [heading, "", *labelled_lines(fields), "", narrative]
    path.write_text("\n".join(content) + "\n", encoding="utf-8")


def write_markdown(path: Path, fields: dict[str, Any], heading: str, narrative: str) -> None:
    content = [f"# {heading}", "", *labelled_lines(fields), "", "## Summary", narrative]
    path.write_text("\n".join(content) + "\n", encoding="utf-8")


def write_json(path: Path, fields: dict[str, Any], narrative: str) -> None:
    payload = dict(fields)
    payload["Summary"] = narrative
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def xml_tag(label: str) -> str:
    return label.lower().replace(" ", "_").replace("-", "_")


def write_xml(path: Path, fields: dict[str, Any], narrative: str, root_name: str) -> None:
    root = ElementTree.Element(root_name)
    for key, value in fields.items():
        child = ElementTree.SubElement(root, xml_tag(key))
        child.text = str(value)
    summary = ElementTree.SubElement(root, "summary")
    summary.text = narrative
    ElementTree.indent(root, space="  ")
    path.write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        + ElementTree.tostring(root, encoding="unicode")
        + "\n",
        encoding="utf-8",
    )


def write_html(path: Path, fields: dict[str, Any], heading: str, narrative: str) -> None:
    rows = "\n".join(
        f"      <p>{key} | {value}</p>" for key, value in fields.items()
    )
    path.write_text(
        "\n".join(
            [
                "<!doctype html>",
                '<html lang="en">',
                "<head>",
                '  <meta charset="utf-8">',
                f"  <title>{heading}</title>",
                "  <style>",
                "    body { font-family: Arial, sans-serif; margin: 40px; color: #17212b; }",
                "    h1 { color: #0f766e; }",
                "    p { margin: 8px 0; }",
                "    .summary { margin-top: 24px; padding: 14px; background: #eef7f5; }",
                "  </style>",
                "</head>",
                "<body>",
                f"  <h1>{heading}</h1>",
                rows,
                f'  <p class="summary">{narrative}</p>',
                "</body>",
                "</html>",
                "",
            ]
        ),
        encoding="utf-8",
    )


def write_docx(path: Path, fields: dict[str, Any], heading: str, narrative: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    title = document.add_heading(heading, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(15, 118, 110)
    subtitle = document.add_paragraph("ProofChain Synthetic Institutional Evidence")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"
    for key, value in fields.items():
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)
    document.add_heading("Summary", level=1)
    document.add_paragraph(narrative)
    footer = section.footer.paragraphs[0]
    footer.text = "Synthetic fixture - no real student or institutional data"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 98, 108)
    document.save(path)


def write_pdf(path: Path, fields: dict[str, Any], heading: str, narrative: str) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "MockTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0F766E"),
        fontSize=18,
        leading=22,
        spaceAfter=12,
    )
    body_style = styles["BodyText"]
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=heading,
        author="ProofChain Synthetic Data Generator",
    )
    story: list[Any] = [
        Paragraph(heading, title_style),
        Paragraph("ProofChain Synthetic Institutional Evidence", styles["Italic"]),
        Spacer(1, 8),
    ]
    table_data = [[Paragraph("<b>Field</b>", body_style), Paragraph("<b>Value</b>", body_style)]]
    table_data.extend(
        [Paragraph(key, body_style), Paragraph(str(value), body_style)]
        for key, value in fields.items()
    )
    table = Table(table_data, colWidths=[58 * mm, 104 * mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#E8F3F1")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#A9B7B5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend(
        [
            table,
            Spacer(1, 14),
            Paragraph("<b>Summary</b>", styles["Heading2"]),
            Paragraph(narrative, body_style),
            Spacer(1, 18),
            Paragraph(
                "Synthetic fixture - no real student or institutional data",
                styles["Italic"],
            ),
        ]
    )
    document.build(story)


def write_structured_document(
    path: Path,
    fields: dict[str, Any],
    heading: str,
    narrative: str,
) -> None:
    extension = path.suffix.lower()
    if extension == ".pdf":
        write_pdf(path, fields, heading, narrative)
    elif extension == ".docx":
        write_docx(path, fields, heading, narrative)
    elif extension == ".html":
        write_html(path, fields, heading, narrative)
    elif extension == ".md":
        write_markdown(path, fields, heading, narrative)
    elif extension == ".json":
        write_json(path, fields, narrative)
    elif extension == ".xml":
        write_xml(path, fields, narrative, "proofchain_evidence")
    elif extension == ".txt":
        write_plain(path, fields, heading, narrative)
    else:
        raise ValueError(f"Unsupported generated document extension: {extension}")


def attendance_rows(fields: dict[str, Any], students: list[Student]) -> list[list[Any]]:
    headers = [
        "Event ID",
        "Event Title",
        "Department",
        "Academic Year",
        "Event Date",
        "Roll Number",
        "Student Name",
        "Gender",
        "Attendance Status",
        "Signature",
    ]
    rows = [headers]
    for student in students:
        rows.append(
            [
                fields["Event ID"],
                fields["Event Title"],
                fields["Department"],
                fields["Academic Year"],
                fields["Event Date"],
                student.roll_number,
                student.full_name,
                student.gender,
                "Present",
                f"Signed-{student.roll_number}",
            ]
        )
    return rows


def write_delimited(path: Path, rows: list[list[Any]], delimiter: str) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter, lineterminator="\n")
        writer.writerows(rows)


def write_student_master_files(root: Path, students_by_department: dict[str, list[Student]]) -> None:
    master = root / "student_master"
    master.mkdir(parents=True, exist_ok=True)
    aiml_payload = {
        "department": "AIML",
        "academic_year": ACADEMIC_YEAR,
        "students": [asdict(student) for student in students_by_department["AIML"]],
    }
    (master / "AIML_students.json").write_text(
        json.dumps(aiml_payload, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )

    aids_root = ElementTree.Element(
        "student_master",
        {"department": "AIDS", "academic_year": ACADEMIC_YEAR},
    )
    for student in students_by_department["AIDS"]:
        node = ElementTree.SubElement(aids_root, "student")
        for key, value in asdict(student).items():
            field = ElementTree.SubElement(node, key)
            field.text = str(value).lower() if isinstance(value, bool) else str(value)
    ElementTree.indent(aids_root, space="  ")
    (master / "AIDS_students.xml").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        + ElementTree.tostring(aids_root, encoding="unicode")
        + "\n",
        encoding="utf-8",
    )

    cse_path = master / "CSE_students.csv"
    fieldnames = list(asdict(students_by_department["CSE"][0]).keys())
    with cse_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        writer.writerows(asdict(student) for student in students_by_department["CSE"])


def write_photo(path: Path, fields: dict[str, Any]) -> None:
    image = Image.new("RGB", (1280, 720), "#E8F3F1")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.load_default(size=38)
    body_font = ImageFont.load_default(size=24)
    small_font = ImageFont.load_default(size=18)
    draw.rectangle((0, 0, 1280, 92), fill="#0F766E")
    draw.text((48, 24), fields["Department"], fill="white", font=title_font)
    draw.text((48, 145), fields["Event Title"], fill="#17212B", font=title_font)
    draw.text((48, 240), fields["Event ID"], fill="#0F766E", font=body_font)
    draw.text((48, 290), fields["Event Date"], fill="#334155", font=body_font)
    draw.rounded_rectangle((48, 390, 1230, 600), radius=16, fill="#FFFFFF", outline="#8FB7B1", width=3)
    draw.text(
        (82, 435),
        "Synthetic mock photo evidence",
        fill="#17212B",
        font=title_font,
    )
    draw.text(
        (82, 505),
        "Generated for deterministic ProofChain testing. No real people are shown.",
        fill="#475569",
        font=small_font,
    )
    image.save(path, format="PNG", optimize=True)


def prepare(root: Path) -> None:
    preserved_readme = (
        (root / "README.md").read_text(encoding="utf-8")
        if (root / "README.md").exists()
        else None
    )
    if root.exists():
        resolved = root.resolve()
        if resolved.name != "mock_institution" or resolved.parent.name != "sample_data":
            raise ValueError(f"Refusing to clean unexpected target: {resolved}")
        shutil.rmtree(root)
    root.mkdir(parents=True)
    if preserved_readme:
        (root / "README.md").write_text(preserved_readme, encoding="utf-8")

    students_by_department = {
        department: build_students(department, index)
        for index, department in enumerate(DEPARTMENTS)
    }
    write_student_master_files(root, students_by_department)

    workbook_payload: dict[str, Any] = {
        "academic_year": ACADEMIC_YEAR,
        "student_master_output": str(root / "student_master" / "all_students.xlsx"),
        "students": [
            asdict(student)
            for department in DEPARTMENTS
            for student in students_by_department[department]
        ],
        "attendance_workbooks": [],
    }
    generation_events: list[dict[str, Any]] = []

    for department in DEPARTMENTS:
        department_root = root / "departments" / department
        for folder in (
            "event_reports",
            "attendance_sheets",
            "approval_documents",
            "certificates",
            "photos",
        ):
            (department_root / folder).mkdir(parents=True, exist_ok=True)

        students = students_by_department[department]
        for index, requirement in enumerate(REQUIREMENTS):
            fields = event_fields(department, requirement, index)
            event_id = fields["Event ID"]
            narrative = (
                f"This synthetic {REQUIREMENTS[requirement].lower()} record documents "
                f"a governed activity for {department}. All 30 registered students are "
                "listed once in the linked attendance evidence, and all approval and "
                "completion records use the same event identity."
            )

            report_extension = EVENT_REPORT_FORMATS[department][index]
            report_path = (
                department_root
                / "event_reports"
                / f"{event_id}_{department}_{requirement}_event_report.{report_extension}"
            )
            write_structured_document(
                report_path,
                fields,
                "Event Report",
                narrative,
            )

            attendance_extension = ATTENDANCE_FORMATS[department][index]
            attendance_path = (
                department_root
                / "attendance_sheets"
                / f"{event_id}_{department}_{requirement}_attendance.{attendance_extension}"
            )
            rows = attendance_rows(fields, students)
            if attendance_extension == "xlsx":
                workbook_payload["attendance_workbooks"].append(
                    {
                        "output": str(attendance_path),
                        "event": fields,
                        "rows": rows,
                    }
                )
            elif attendance_extension == "csv":
                write_delimited(attendance_path, rows, ",")
            else:
                write_delimited(attendance_path, rows, "\t")

            approval_extension = APPROVAL_FORMATS[department][index]
            approval_fields = dict(fields)
            approval_fields.update(
                {
                    "Approval Status": "Approved",
                    "Signature Present": "Yes",
                    "Approval Number": f"APR-{department}-{index + 1:03d}-2025",
                    "Reference Number": f"REF-{event_id}",
                }
            )
            approval_path = (
                department_root
                / "approval_documents"
                / f"APR_{event_id}_{department}_{requirement}_approval.{approval_extension}"
            )
            write_structured_document(
                approval_path,
                approval_fields,
                "Approval Document",
                "The Head of Department approved this synthetic activity and its evidence plan.",
            )

            certificate_extension = CERTIFICATE_FORMATS[department][index]
            certificate_fields = dict(fields)
            certificate_fields.update(
                {
                    "Certificate Register": "Complete",
                    "Awarded To": "All 30 registered participants",
                    "Signature Present": "Yes",
                    "Reference Number": f"CERT-{event_id}",
                }
            )
            certificate_path = (
                department_root
                / "certificates"
                / f"CERT_{event_id}_{department}_{requirement}_certificate.{certificate_extension}"
            )
            write_structured_document(
                certificate_path,
                certificate_fields,
                "Certificate of Participation",
                "This synthetic register confirms completion for the 30 linked attendance records.",
            )

            photo_path = department_root / "photos" / f"PHOTO_{event_id}_{requirement}.png"
            write_photo(photo_path, fields)
            generation_events.append(
                {
                    "event_id": event_id,
                    "department": department,
                    "requirement_id": requirement,
                    "event_title": fields["Event Title"],
                    "event_date": fields["Event Date"],
                    "reported_participant_count": 30,
                    "unique_attendance_count": 30,
                    "expected_status": "defensible",
                    "files": {
                        "event_report": str(report_path.relative_to(root)),
                        "attendance_sheet": str(attendance_path.relative_to(root)),
                        "approval_document": str(approval_path.relative_to(root)),
                        "certificate": str(certificate_path.relative_to(root)),
                        "photograph": str(photo_path.relative_to(root)),
                    },
                }
            )

    (root / "workbook_payload.json").write_text(
        json.dumps(workbook_payload, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )
    (root / "generation_events.json").write_text(
        json.dumps(generation_events, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def finalize(root: Path) -> None:
    payload = json.loads((root / "workbook_payload.json").read_text(encoding="utf-8"))
    missing_workbooks = [
        path
        for path in [
            Path(payload["student_master_output"]),
            *(Path(item["output"]) for item in payload["attendance_workbooks"]),
        ]
        if not path.exists()
    ]
    if missing_workbooks:
        raise FileNotFoundError(f"Workbook builder has not created: {missing_workbooks}")

    events = json.loads((root / "generation_events.json").read_text(encoding="utf-8"))
    evidence_files = sorted(
        path
        for path in (root / "departments").rglob("*")
        if path.is_file()
    )
    student_files = sorted(
        path for path in (root / "student_master").rglob("*") if path.is_file()
    )
    formats: dict[str, int] = {}
    for path in [*evidence_files, *student_files]:
        formats[path.suffix.lower()] = formats.get(path.suffix.lower(), 0) + 1
    departments = {}
    for department in DEPARTMENTS:
        department_students = [
            student for student in payload["students"] if student["department"] == department
        ]
        departments[department] = {
            "name": DEPARTMENTS[department]["name"],
            "student_count": len(department_students),
            "female_count": sum(
                student["gender"] == "Female" for student in department_students
            ),
            "male_count": sum(
                student["gender"] == "Male" for student in department_students
            ),
            "event_count": sum(event["department"] == department for event in events),
        }
    manifest = {
        "dataset_id": "PROOFCHAIN-MOCK-THREE-DEPARTMENT-2025-2026",
        "schema_version": "1.0.0",
        "synthetic": True,
        "academic_year": ACADEMIC_YEAR,
        "source_root": str((root / "departments").resolve()),
        "department_count": len(departments),
        "student_count": len(payload["students"]),
        "female_count": sum(student["gender"] == "Female" for student in payload["students"]),
        "male_count": sum(student["gender"] == "Male" for student in payload["students"]),
        "departments": departments,
        "requirements": REQUIREMENTS,
        "event_count": len(events),
        "evidence_file_count": len(evidence_files),
        "supported_format_inventory": dict(sorted(formats.items())),
        "events": events,
        "files": [
            {
                "relative_path": str(path.relative_to(root)).replace("\\", "/"),
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in [*student_files, *evidence_files]
        ],
    }
    (root / "dataset_manifest.json").write_text(
        json.dumps(manifest, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )
    expected = {
        "department_count": 3,
        "student_count": 90,
        "students_per_department": 30,
        "gender_distribution_per_department": {"Female": 15, "Male": 15},
        "event_count": 15,
        "evidence_file_count": 75,
        "expected_event_bundle_count": 15,
        "expected_missing_required_documents": 0,
        "expected_participant_count_mismatches": 0,
        "expected_duplicate_roll_numbers": 0,
        "expected_unsigned_approvals": 0,
        "expected_unsupported_files": 0,
        "expected_rejected_files": 0,
        "notes": [
            "Photographs use metadata-only processing and are optional evidence.",
            "All people, emails, approvals, signatures, and events are synthetic.",
            "The external-submission adapter remains governed by environment credentials.",
        ],
    }
    (root / "expected_outcomes.json").write_text(
        json.dumps(expected, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("action", choices=["prepare", "finalize"])
    parser.add_argument(
        "--root",
        type=Path,
        default=Path("sample_data/mock_institution"),
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.action == "prepare":
        prepare(args.root)
    else:
        finalize(args.root)


if __name__ == "__main__":
    main()
